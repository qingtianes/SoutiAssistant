# -*- coding: utf-8 -*-
"""
搜题助手 · docx 题库切块验证工具
用户方案：docx 只做"切块"，不拆字段。题干+选项+答案整块原样保留。
验证：切块成功率、漏切/错切统计。

用法：
    python docx_chunk_verify.py <docx文件或目录>
"""

import sys
import os
import re
import glob


# 章节标题判断：一、单选题 / 二、多选题 / 三、判断题 / 动火作业 / 1.动火作业 / 临时用电作业
def is_section_title(t):
    if not t:
        return True
    # "一、单选题" / "二、多选题。" / "三、判断题。"
    if re.match(r'^[一二三四五六七八九十]+、', t):
        return True
    # 纯章节名（无数字开头、无括号、含"作业"且短）
    if '作业' in t and len(t) < 10 and not re.search(r'[（(]', t) and not re.match(r'^\d+[.、]', t):
        return True
    # 数字开头的短标题：1.动火作业 / 2.吊装作业
    if re.match(r'^\d+[.、]\S{1,6}$', t) and not re.search(r'[（(]', t):
        return True
    # 页眉/页脚/空说明：如 "1/105" "2/105" "一、判断题，共206道。"
    if re.match(r'^\d+/\d+$', t):
        return True
    if re.match(r'^[一二三四五六七八九十]+、.+共\d+道', t):
        return True
    return False


def is_question_start(t):
    """题目开始：数字开头 + （有括号或足够长）"""
    m = re.match(r'^(\d+)[.、]\s*(.+)$', t)
    if not m:
        return False
    stem = m.group(2).strip()
    if len(stem) < 4 and not re.search(r'[（(]', stem):
        return False  # 太短，可能是标题
    return True


def chunk_docx(path):
    """把 docx 切成题目块，返回 [(题号, 块文本), ...] 和统计"""
    import docx
    doc = docx.Document(path)
    txts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    chunks = []  # (题号, [行列表])
    stats = {'skipped_sections': 0, 'orphan_lines': []}
    current = None
    current_qid = None

    for t in txts:
        if is_section_title(t):
            stats['skipped_sections'] += 1
            continue
        if is_question_start(t):
            # 新题目开始，保存上一块
            if current is not None:
                chunks.append((current_qid, current))
            m = re.match(r'^(\d+)[.、]\s*(.+)$', t)
            current_qid = int(m.group(1))
            current = [t]
        else:
            # 非标题、非题目开始的行 → 归属当前块（选项/答案/续行）
            if current is not None:
                current.append(t)
            else:
                stats['orphan_lines'].append(t)

    if current is not None:
        chunks.append((current_qid, current))

    return chunks, stats


def verify(path):
    import docx
    doc = docx.Document(path)
    txts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    chunks, stats = chunk_docx(path)

    # 原始题干行数（用于对比漏切）
    all_q = []
    for t in txts:
        if is_question_start(t):
            m = re.match(r'^(\d+)[.、]', t)
            all_q.append(int(m.group(1)))

    chunk_qids = [qid for qid, _ in chunks]

    # 漏切：原题干有但切块里没有
    missing = [q for q in all_q if q not in chunk_qids]

    print(f'文件: {os.path.basename(path)}')
    print(f'  原题干行数: {len(all_q)}  切块数: {len(chunks)}')
    print(f'  漏切题号: {missing if missing else "无"}')
    if stats['orphan_lines']:
        print(f'  孤儿行（未归属任何块）: {len(stats["orphan_lines"])} 行')
        for o in stats['orphan_lines'][:5]:
            print(f'    - {o[:60]}')
    print()

    # 抽查第一块和最后一块
    if chunks:
        print('  ── 切块样例 ──')
        for qid, lines in chunks[:2]:
            print(f'  [块 {qid}]')
            for l in lines:
                print(f'    {l[:70]}')
        print('  ...')
    return len(all_q), len(chunks), len(missing)


def main():
    if len(sys.argv) < 2:
        print(__doc__)
        sys.exit(1)
    src = sys.argv[1]
    files = []
    if os.path.isdir(src):
        for ext in ('.docx',):
            files += glob.glob(os.path.join(src, '**', '*' + ext), recursive=True)
    else:
        files = [src]

    total_q = total_chunk = total_miss = 0
    for f in sorted(files):
        if os.path.basename(f).startswith('~$'):
            continue
        try:
            q, c, m = verify(f)
            total_q += q
            total_chunk += c
            total_miss += m
        except Exception as e:
            print(f'❌ {os.path.basename(f)}: {e}')

    print(f'\n===== 汇总 =====')
    print(f'总题干: {total_q}  总切块: {total_chunk}  总漏切: {total_miss}')


if __name__ == '__main__':
    main()
